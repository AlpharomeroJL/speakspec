"""Export writers (Phase 9): the full artifact bundle on disk.

Writes into a session directory:

* ``AGENTS.md`` + ``CLAUDE.md`` (through the Phase 6 quality gate),
* ``docs/adr/NNNN-*.md`` (one MADR file per decision),
* ``diagrams/*.mmd`` (raw Mermaid source),
* ``spec.md`` (full Markdown export, diagrams inline),
* ``spec.docx`` (formatted Word export),
* ``spec.json`` (full structured spec for programmatic use).
"""

import json
from pathlib import Path
from typing import Any

from speakspec.agents_md import write_agent_files
from speakspec.mermaid_repair import DIAGRAM_KINDS
from speakspec.messages import SidecarError
from speakspec.rpc import RequestContext

DIAGRAM_TITLES = {
    "sequence": "Sequence — main user flow",
    "er": "Entity-Relationship — data model",
    "component": "Component — boundaries and data flow",
    "c4_context": "C4 L1 — system context",
    "c4_container": "C4 L2 — containers",
}


def render_markdown(stage2: dict[str, Any], stage3: dict[str, Any]) -> str:
    """Compose the full Markdown export with diagrams inline."""
    lines: list[str] = []
    add = lines.append
    add("# Architecture spec\n")
    add(stage2.get("system_overview", ""))

    lang = stage2.get("recommended_language", {})
    add("\n## Recommended language\n")
    add(
        f"**{lang.get('language', '?')}** — confidence {lang.get('confidence', '?')}"
        + (" (user override)" if lang.get("overridden_by_user") else "")
    )
    add(f"\n{lang.get('rationale', '')}")

    add("\n## Quality goals\n")
    for goal in stage2.get("quality_goals", []):
        add(f"- **{goal.get('goal', '')}** — driven by: {goal.get('driver', '')}")

    add("\n## Architecture decisions\n")
    for i, decision in enumerate(stage2.get("architecture_decisions", []), start=1):
        add(f"### {i}. {decision.get('title', '')}")
        add(
            f"**Chosen:** {decision.get('chosen_option', '')} "
            f"(confidence: {decision.get('confidence', '?')})"
        )
        add(f"\n{decision.get('rationale', '')}\n")
        consequences = decision.get("consequences", {})
        for positive in consequences.get("positive", []):
            add(f"- ✅ {positive}")
        for negative in consequences.get("negative", []):
            add(f"- ⚠️ {negative}")
        add("\n**Ruled out:**")
        for ruled in decision.get("ruled_out", []):
            add(f"- {ruled.get('option', '')} — {ruled.get('rejection_reason', '')}")
        add("")

    add("## Use, don't build\n")
    for component in stage2.get("oss_components", []):
        add(
            f"- **{component.get('name', '')}** {component.get('version', '')} — "
            f"{component.get('purpose', '')}. {component.get('why_selected', '')}"
        )
    add("\n## What not to build\n")
    for item in stage2.get("what_not_to_build", []):
        add(f"- {item}")

    add("\n## Risk flags\n")
    for risk in stage2.get("risk_flags", []):
        add(
            f"- **[{risk.get('severity', '?')}] {risk.get('risk', '')}** — "
            f"mitigation: {risk.get('mitigation', '')}"
        )

    add("\n## Open questions\n")
    for question in stage2.get("open_questions", []):
        add(f"- {question}")

    complexity = stage2.get("complexity", {})
    add("\n## Complexity\n")
    add(
        f"Rating **{complexity.get('rating', '?')}** — {complexity.get('time_estimate', '')}. "
        f"Key drivers: {', '.join(complexity.get('key_drivers', []))}"
    )

    add("\n## Hard constraints\n")
    for constraint in stage2.get("constraints", []):
        add(f"- {constraint}")

    add("\n## Diagrams\n")
    diagrams = stage3.get("diagrams", {})
    for kind in DIAGRAM_KINDS:
        add(f"### {DIAGRAM_TITLES[kind]}\n")
        add("```mermaid")
        add(diagrams.get(kind, ""))
        add("```\n")

    add("## Suggested file tree\n")
    add("```")
    add(stage3.get("file_tree", ""))
    add("```\n")

    first_pr = stage3.get("first_pr", {})
    add("## First PR\n")
    add(f"**{first_pr.get('title', '')}** — {first_pr.get('scope', '')}")
    add(f"\nEstimate: {first_pr.get('time_estimate', '')}\n")
    for task in first_pr.get("tasks", []):
        add(f"1. {task}")
    out_of_scope = first_pr.get("out_of_scope", [])
    if out_of_scope:
        add("\nOut of scope: " + "; ".join(out_of_scope))
    add("")
    return "\n".join(lines)


def render_docx(stage2: dict[str, Any], stage3: dict[str, Any], dest: Path) -> None:
    """Write the formatted Word export (stakeholder sharing)."""
    import docx
    from docx.shared import Pt

    document = docx.Document()
    document.add_heading("Architecture spec", level=0)
    document.add_paragraph(stage2.get("system_overview", ""))

    lang = stage2.get("recommended_language", {})
    document.add_heading("Recommended language", level=1)
    document.add_paragraph(
        f"{lang.get('language', '?')} — confidence {lang.get('confidence', '?')}"
        + (" (user override)" if lang.get("overridden_by_user") else "")
    )
    document.add_paragraph(lang.get("rationale", ""))

    document.add_heading("Architecture decisions", level=1)
    for decision in stage2.get("architecture_decisions", []):
        document.add_heading(decision.get("title", ""), level=2)
        document.add_paragraph(
            f"Chosen: {decision.get('chosen_option', '')} "
            f"(confidence: {decision.get('confidence', '?')})"
        )
        document.add_paragraph(decision.get("rationale", ""))
        for ruled in decision.get("ruled_out", []):
            document.add_paragraph(
                f"Ruled out: {ruled.get('option', '')} — {ruled.get('rejection_reason', '')}",
                style="List Bullet",
            )

    document.add_heading("Risks", level=1)
    for risk in stage2.get("risk_flags", []):
        document.add_paragraph(
            f"[{risk.get('severity', '?')}] {risk.get('risk', '')} — {risk.get('mitigation', '')}",
            style="List Bullet",
        )

    document.add_heading("Diagrams (Mermaid source)", level=1)
    diagrams = stage3.get("diagrams", {})
    for kind in DIAGRAM_KINDS:
        document.add_heading(DIAGRAM_TITLES[kind], level=2)
        paragraph = document.add_paragraph(diagrams.get(kind, ""))
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    document.add_heading("Suggested file tree", level=1)
    tree_paragraph = document.add_paragraph(stage3.get("file_tree", ""))
    for run in tree_paragraph.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    document.save(str(dest))


def write_bundle(
    session_dir: Path,
    stage2: dict[str, Any],
    stage3: dict[str, Any],
    diagram_reports: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Write every export artifact; returns written paths + the agents gate report."""
    session_dir.mkdir(parents=True, exist_ok=True)
    written: list[str] = []

    gate_report = write_agent_files(
        stage3.get("agents_md", ""),
        stage3.get("claude_md_shim", ""),
        stage3.get("file_tree", ""),
        session_dir,
    )
    written.extend(gate_report.pop("written"))

    adr_dir = session_dir / "docs" / "adr"
    adr_dir.mkdir(parents=True, exist_ok=True)
    for adr in stage3.get("adrs", []):
        # Filenames are schema-constrained to docs/adr/NNNN-slug.md.
        name = Path(adr.get("filename", "docs/adr/0000-unnamed.md")).name
        path = adr_dir / name
        path.write_text(adr.get("markdown", ""), encoding="utf-8", newline="\n")
        written.append(str(path))

    diagram_dir = session_dir / "diagrams"
    diagram_dir.mkdir(parents=True, exist_ok=True)
    for kind, source in stage3.get("diagrams", {}).items():
        path = diagram_dir / f"{kind}.mmd"
        path.write_text(source, encoding="utf-8", newline="\n")
        written.append(str(path))

    markdown_path = session_dir / "spec.md"
    markdown_path.write_text(render_markdown(stage2, stage3), encoding="utf-8", newline="\n")
    written.append(str(markdown_path))

    docx_path = session_dir / "spec.docx"
    render_docx(stage2, stage3, docx_path)
    written.append(str(docx_path))

    json_path = session_dir / "spec.json"
    json_path.write_text(
        json.dumps(
            {
                "architecture_spec": stage2,
                "output_bundle": stage3,
                "diagram_reports": diagram_reports or [],
            },
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
        newline="\n",
    )
    written.append(str(json_path))

    return {"written": written, "agents_gate": gate_report}


def handle_export_bundle(params: dict[str, Any], ctx: RequestContext) -> dict[str, Any]:
    """RPC handler: write the full bundle for a session.

    Params: ``session_dir``, ``architecture_spec`` (Stage 2),
    ``output_bundle`` (Stage 3), optional ``diagram_reports``.
    """
    session_dir = params.get("session_dir", "")
    stage2 = params.get("architecture_spec")
    stage3 = params.get("output_bundle")
    if not session_dir or not stage2 or not stage3:
        raise SidecarError(
            "missing-params",
            "export.bundle needs session_dir, architecture_spec, and output_bundle.",
        )
    ctx.emit_progress({"state": "writing-exports"})
    return write_bundle(Path(session_dir), stage2, stage3, params.get("diagram_reports"))
